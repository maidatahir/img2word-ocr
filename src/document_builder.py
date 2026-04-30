"""
Document Builder Module
Converts structured OCR data into professional document formats 
including Microsoft Word (.docx), PDF, and Plain Text.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

# Document Styling Constants
bodyFont      = "Calibri"
headingFont   = "Calibri"
bodyFontSize  = 11
heading1Size  = 14
heading2Size  = 12

# Map internal alignment strings to python-docx alignment constants
alignMap = {
    "LEFT":   WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT":  WD_ALIGN_PARAGRAPH.RIGHT,
}


def applyParagraphFormat(para, alignValue: str):
    """Configures spacing and alignment for a Word paragraph."""
    para.alignment = alignMap.get(alignValue, WD_ALIGN_PARAGRAPH.LEFT)
    paraFmt = para.paragraph_format
    paraFmt.space_before = Pt(4)
    paraFmt.space_after  = Pt(4)
    paraFmt.line_spacing = Pt(14)


def addHeading(doc: Document, headingText: str, headingLevel: int = 1):
    """Adds a formatted heading to the Word document."""
    cleanHeading = headingText.lstrip("#").strip()
    headingPara = doc.add_heading(cleanHeading, level=headingLevel)
    headingPara.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    for headingRun in headingPara.runs:
        headingRun.font.name      = headingFont
        headingRun.font.size      = Pt(heading1Size if headingLevel == 1 else heading2Size)
        headingRun.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E) # Professional dark blue/black
    return headingPara


def addBodyParagraph(doc: Document, block: dict):
    """Adds a standard body paragraph with potential bold/italic formatting."""
    bodyPara = doc.add_paragraph()
    applyParagraphFormat(bodyPara, block.get("alignment", "LEFT"))
    
    runsList = block.get("runs", [])
    if runsList:
        # Construct paragraph word by word to apply inline formatting
        for runData in runsList:
            docRun = bodyPara.add_run(runData.get("text", ""))
            docRun.bold        = runData.get("bold", False)
            docRun.italic      = runData.get("italic", False)
            docRun.font.name   = bodyFont
            docRun.font.size   = Pt(bodyFontSize)
    else:
        # Fallback for plain text blocks
        plainRun = bodyPara.add_run(block.get("text", ""))
        plainRun.font.name = bodyFont
        plainRun.font.size = Pt(bodyFontSize)
    return bodyPara


def buildDocx(ocrResult: dict, outputPath: str) -> str:
    """Generates a Microsoft Word document from OCR results."""
    doc = Document()
    
    # Configure standard academic margins
    for docSection in doc.sections:
        docSection.top_margin    = Inches(1.0)
        docSection.bottom_margin = Inches(1.0)
        docSection.left_margin   = Inches(1.25)
        docSection.right_margin  = Inches(1.25)
        
    normalStyle = doc.styles["Normal"]
    normalStyle.font.name = bodyFont
    normalStyle.font.size = Pt(bodyFontSize)
    
    paragraphList = ocrResult.get("paragraphs", [])
    if not paragraphList:
        # Handle raw text if structural paragraph data is missing
        rawContent = ocrResult.get("rawText", "")
        for rawLine in rawContent.splitlines():
            rawLine = rawLine.strip()
            if not rawLine:
                doc.add_paragraph()
                continue
            fallbackPara = doc.add_paragraph()
            fallbackRun = fallbackPara.add_run(rawLine)
            fallbackRun.font.name = bodyFont
            fallbackRun.font.size = Pt(bodyFontSize)
    else:
        # Process structural blocks (headings and body paragraphs)
        headingCount = 0
        for block in paragraphList:
            blockText = block.get("text", "").strip()
            if not blockText:
                doc.add_paragraph()
                continue
            if block.get("isHeading"):
                headingCount += 1
                headingLevel = 1 if headingCount == 1 else 2
                addHeading(doc, blockText, headingLevel=headingLevel)
            else:
                addBodyParagraph(doc, block)
                
    # Ensure directory exists and save the file
    os.makedirs(os.path.dirname(os.path.abspath(outputPath)), exist_ok=True)
    doc.save(outputPath)
    return outputPath


def buildPdf(ocrResult: dict, outputPath: str) -> str:
    """Generates a basic PDF document from OCR results."""
    pdfCanvas = canvas.Canvas(outputPath, pagesize=letter)
    pageWidth, pageHeight = letter
    yPosition = pageHeight - 1*inch
    
    rawText = ocrResult.get("rawText", "")
    if not rawText:
        lines = [p.get("text", "") for p in ocrResult.get("paragraphs", [])]
        rawText = "\n".join(lines)
        
    pdfCanvas.setFont("Helvetica", 10)
    for line in rawText.splitlines():
        # Handle page breaks
        if yPosition < 1*inch:
            pdfCanvas.showPage()
            yPosition = pageHeight - 1*inch
            pdfCanvas.setFont("Helvetica", 10)
            
        pdfCanvas.drawString(1*inch, yPosition, line)
        yPosition -= 15
        
    pdfCanvas.save()
    return outputPath


def buildTxt(ocrResult: dict, outputPath: str) -> str:
    """Saves OCR results as a plain text file."""
    rawText = ocrResult.get("rawText", "")
    if not rawText:
        lines = [p.get("text", "") for p in ocrResult.get("paragraphs", [])]
        rawText = "\n".join(lines)
        
    with open(outputPath, "w", encoding="utf-8") as f:
        f.write(rawText)
    return outputPath
